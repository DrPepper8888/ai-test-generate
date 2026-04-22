"""
Create P1→P2 Development Word Document
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PROJECT_ROOT = Path("D:/pejic/AAAHarnessEngineering/ai-test-case-generator")

doc = Document()

# Title
title = doc.add_heading("P1 → P2 开发文档", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1
doc.add_heading("项目状态", level=1)
doc.add_paragraph("P0：✅ 已完成（最小可用版本）")
doc.add_paragraph("P1：✅ 已完成（质量优化）")
doc.add_paragraph("P2：✅ 已完成（记忆与学习系统）")

# Section 2
doc.add_heading("P2 开发内容总览", level=1)
doc.add_paragraph("核心目标：从\"质量优化\"升级到\"学习优化\"，让 AI 能够从用户的反馈中学习，越用越准。")

# Section 3
doc.add_heading("P2 新增功能（详细）", level=1)

doc.add_heading("1. 会话持久化系统", level=2)
doc.add_paragraph("实现思路：")
p = doc.add_paragraph("之前：会话只在内存，刷新页面丢失")
p.style = "List Bullet"
p = doc.add_paragraph("现在：会话保存到 data/memory/sessions/{session_id}.json")
p.style = "List Bullet"
p = doc.add_paragraph("URL 自动带 ?session=xxx，刷新不丢失")
p.style = "List Bullet"

doc.add_heading("2. 用例标签系统", level=2)
doc.add_paragraph("功能说明：")
p = doc.add_paragraph("✅ 可投产 — 标记高质量用例")
p.style = "List Bullet"
p = doc.add_paragraph("❌ 舍弃 — 标记低质量用例")
p.style = "List Bullet"
p = doc.add_paragraph("⚠️ 待修改 — 标记需要改进的用例")
p.style = "List Bullet"

doc.add_heading("3. 用例编辑功能", level=2)
doc.add_paragraph("功能说明：")
p = doc.add_paragraph("支持手动编辑任意用例的任意字段")
p.style = "List Bullet"
p = doc.add_paragraph("保存编辑前后对照")
p.style = "List Bullet"
p = doc.add_paragraph("显示\"✏️已编辑\"标记")
p.style = "List Bullet"

doc.add_heading("4. 多轮对话反馈", level=2)
doc.add_paragraph("功能说明：")
p = doc.add_paragraph("用户输入反馈内容")
p.style = "List Bullet"
p = doc.add_paragraph("AI 只修改该条用例")
p.style = "List Bullet"
p = doc.add_paragraph("其他用例完全不变")
p.style = "List Bullet"

doc.add_heading("5. 规则学习系统", level=2)
doc.add_paragraph("核心设计，借鉴 OpenClaw 的记忆与学习机制：")
doc.add_paragraph("1. 规则提取：用户标记用例后，点击\"📚学习规则\"，AI 自动从标签和反馈中提取规则，保存到 data/memory/rules.json")
doc.add_paragraph("2. 规则注入：下次生成用例时，自动将已学规则注入到 system_prompt，提升生成质量")

doc.add_heading("6. 导出功能增强", level=2)
p = doc.add_paragraph("\"下载选中用例\" → \"下载可投产用例\"")
p.style = "List Bullet"
p = doc.add_paragraph("只下载标为\"✅可投产\"的用例")
p.style = "List Bullet"

# Section 4
doc.add_heading("文件变更清单", level=1)

doc.add_heading("新增文件（P2）", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "文件"
table.rows[0].cells[1].text = "用途"

new_files = [
    ("src/memory/__init__.py", "memory 模块入口"),
    ("src/memory/memory_store.py", "Memory 存储管理"),
    ("src/memory/rule_extractor.py", "规则提取器"),
    ("src/memory/rule_injector.py", "规则注入器"),
    ("项目功能说明书.md", "功能总说明"),
    ("项目依赖说明.md", "依赖说明"),
    ("项目文件结构.md", "文件结构说明"),
    ("验收测试清单.md", "验收测试清单"),
    ("test_e2e.py", "自动化 API 测试"),
    ("test_unit_final.py", "最终单元测试"),
]

for file, desc in new_files:
    row = table.add_row()
    row.cells[0].text = file
    row.cells[1].text = desc

doc.add_heading("修改文件（P2）", level=2)
table2 = doc.add_table(rows=1, cols=2)
table2.style = "Light Grid Accent 1"
table2.rows[0].cells[0].text = "文件"
table2.rows[0].cells[1].text = "修改内容"

mod_files = [
    ("web/app.py", "会话持久化、标签、编辑、反馈、规则学习 API"),
    ("src/workflow/pipeline.py", "支持规则注入、MemoryStore 初始化"),
]

for file, desc in mod_files:
    row = table2.add_row()
    row.cells[0].text = file
    row.cells[1].text = desc

# Section 5
doc.add_heading("总结", level=1)
doc.add_paragraph("P2 阶段完成了从\"质量优化\"到\"学习优化\"的升级，核心新增：")
p = doc.add_paragraph("1. 记忆系统（会话持久化）")
p.style = "List Bullet"
p = doc.add_paragraph("2. 规则学习系统（从反馈中学习）")
p.style = "List Bullet"
p = doc.add_paragraph("3. 多轮交互（标签、编辑、反馈）")
p.style = "List Bullet"
p = doc.add_paragraph("4. 完整自动化测试覆盖")
p.style = "List Bullet"

doc.add_paragraph("项目已可交付！🎉")

# Save
docx_path = PROJECT_ROOT / "P1到P2开发文档.docx"
doc.save(docx_path)
print(f"Word文档已创建：{docx_path}")

# Also copy to transfer folder
TRANSFER_DIR = PROJECT_ROOT / "传输文件夹"
TRANSFER_DIR.mkdir(parents=True, exist_ok=True)
import shutil
shutil.copy2(docx_path, TRANSFER_DIR / "P1到P2开发文档.docx.txt")
print(f"已复制到传输文件夹：{TRANSFER_DIR / 'P1到P2开发文档.docx.txt'}")
